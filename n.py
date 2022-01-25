import sys, os
import docx
from docx.shared import Inches
def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

# Create an instance of a word document
try:
	doc = docx.Document('commands_git.docx')
except Exception as e:
	doc = docx.Document()
command = sys.argv[1]
image = sys.argv[2]

# Image in its native size
#doc.add_heading('Image in Native Size:', 3)
cleaned_string = ''.join(c for c in command if valid_xml_char_ordinal(c))

doc.add_heading(cleaned_string, 3)
doc.add_picture(image, height=Inches(3.5))

''' Delete '''
if(os.path.isfile(image)):
    os.remove(image)

# Now save the document to a location
doc.save('commands_git.docx')
